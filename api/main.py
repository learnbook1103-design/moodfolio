import json
import os
import re
import requests
from fastapi import FastAPI, HTTPException, Depends, status, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv

# AI 도구
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.messages import HumanMessage

# DB & 보안 도구
from sqlalchemy import create_engine, Column, Integer, String
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session
from passlib.context import CryptContext

# 구글 인증 도구
from google.oauth2 import id_token
from google.auth.transport import requests as google_requests

# 1. 환경 설정
from pathlib import Path

# .env 파일에서 직접 읽기 (load_dotenv 대신)
env_path = Path(__file__).parent / '.env'
GOOGLE_API_KEY = None

if env_path.exists():
    try:
        # utf-8-sig handles BOM automatically
        with open(env_path, 'r', encoding='utf-8-sig') as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith('#'):
                    continue
                if '=' in line:
                    key, value = line.split('=', 1)
                    key = key.strip()
                    value = value.strip()
                    if key == 'GOOGLE_API_KEY':
                        GOOGLE_API_KEY = value
                        print(f"✅ GOOGLE_API_KEY loaded successfully: {GOOGLE_API_KEY[:10]}...")
                        break
    except Exception as e:
        print(f"⚠️ Error reading .env file: {e}")

if not GOOGLE_API_KEY:
    print(f"❌ GOOGLE_API_KEY not found in {env_path}")
    print(f"❌ File exists: {env_path.exists()}")
    if env_path.exists():
        print("❌ File contents (first 5 lines):")
        try:
            with open(env_path, 'r', encoding='utf-8-sig') as f:
                for i, line in enumerate(f):
                    if i < 5:
                        print(f"  Line {i+1}: {line.strip()[:50]}...")
        except:
            pass
    # Fallback: check os.environ just in case
    GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
    if GOOGLE_API_KEY:
         print(f"✅ GOOGLE_API_KEY loaded from environment variable: {GOOGLE_API_KEY[:10]}...")

if not GOOGLE_API_KEY:
    print(f"⚠️ WARNING: GOOGLE_API_KEY not found - AI features will not work")
    print(f"⚠️ Please set GOOGLE_API_KEY environment variable in Vercel")

app = FastAPI()

# CORS 설정 (모든 주소 허용)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/api/health")
@app.get("/api/health")
def health_check():
    db_type = "postgresql" if SQLALCHEMY_DATABASE_URL.startswith("postgresql") else "sqlite"
    return {
        "status": "ok", 
        "message": "Backend is running!",
        "database_type": db_type,
        "supabase_connected": bool(SUPABASE_URL and SUPABASE_DB_PASSWORD)
    }

# Test endpoint to verify backend is working
@app.get("/backend/test")
def test_backend():
    return {"status": "success", "message": "Backend is working!", "timestamp": "2026-01-16"}

# 2. 데이터베이스 설정 (Supabase PostgreSQL)
SUPABASE_URL = os.getenv("NEXT_PUBLIC_SUPABASE_URL", "")
SUPABASE_DB_PASSWORD = os.getenv("SUPABASE_DB_PASSWORD", "")

import urllib.parse

# Supabase PostgreSQL 연결 문자열 구성
# URL 형식: postgresql://postgres.{project-ref}:{password}@aws-0-ap-northeast-2.pooler.supabase.com:6543/postgres
if SUPABASE_URL and SUPABASE_DB_PASSWORD:
    # Extract project ref safely (remove trailing slashes and whitespace/newlines)
    # Error log showed \r\n in hostname, so we must strip() whitespace first
    project_ref = SUPABASE_URL.replace("https://", "").replace(".supabase.co", "").strip().strip("/")
    
    # URL encode password to handle special characters safely
    encoded_password = urllib.parse.quote_plus(SUPABASE_DB_PASSWORD.strip())
    
    # Use Supavisor Pooler (Transaction Mode) which supports IPv4
    # The direct db. domain resolves to IPv6 which Vercel/Lambda may not support
    # "Tenant not found" error previously was likely due to the trailing \r\n in project_ref
    SQLALCHEMY_DATABASE_URL = f"postgresql://postgres.{project_ref}:{encoded_password}@aws-0-ap-northeast-2.pooler.supabase.com:6543/postgres"
    print(f"📂 Using Supabase PostgreSQL (Pooler): {project_ref}")
else:
    # Fallback to SQLite for local development
    print("⚠️ Supabase credentials not found, falling back to SQLite")
    if os.environ.get("VERCEL") or os.environ.get("AWS_LAMBDA_FUNCTION_NAME"):
        SQLALCHEMY_DATABASE_URL = "sqlite:////tmp/users.db"
        print("📂 Using /tmp/users.db for serverless storage")
    else:
        SQLALCHEMY_DATABASE_URL = "sqlite:///./users.db"

# PostgreSQL은 check_same_thread 옵션이 필요 없음
if SQLALCHEMY_DATABASE_URL.startswith("postgresql"):
    engine = create_engine(SQLALCHEMY_DATABASE_URL, pool_pre_ping=True)
else:
    engine = create_engine(SQLALCHEMY_DATABASE_URL, connect_args={"check_same_thread": False})

SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

# User 테이블 정의
class User(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True, index=True)
    email = Column(String, unique=True, index=True)
    password = Column(String)
    name = Column(String)
    portfolio_data = Column(String, nullable=True)

try:
    Base.metadata.create_all(bind=engine)
    print("✅ Database tables created/verified")
except Exception as e:
    print(f"⚠️ Database initialization failed (Non-critical for Admin APIs): {e}")
    # We continue running because Admin APIs use Supabase HTTP Client, not this SQLAlchemy connection


# 비밀번호 암호화
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# --- 데이터 모델 정의 ---
class UserCreate(BaseModel):
    email: str
    password: str
    name: str

class UserLogin(BaseModel):
    email: str
    password: str

class GoogleToken(BaseModel):
    token: str

class KakaoToken(BaseModel):
    token: str

class NaverToken(BaseModel):
    token: str

class UserAnswers(BaseModel):
    answers: dict

class ChatRequest(BaseModel):
    message: str
    portfolio_context: str | None = None
    is_shared: bool = False

class PortfolioUpdate(BaseModel):
    email: str
    portfolio_data: dict

class ChatAnswerGenerationRequest(BaseModel):
    portfolio_context: str

# --- LLM 초기화 (모든 엔드포인트에서 사용) ---
llm = None
if GOOGLE_API_KEY:
    try:
        llm = ChatGoogleGenerativeAI(
            model="gemini-flash-latest",
            temperature=0.7,
            google_api_key=GOOGLE_API_KEY
        )
        print("✅ LLM initialized successfully")
    except Exception as e:
        print(f"❌ Failed to initialize LLM: {e}")
else:
    print("⚠️ LLM not initialized - GOOGLE_API_KEY missing")

# --- [API] AI 채팅 답변 생성 ---
@app.post("/generate-chat-answers")
def generate_chat_answers(request: ChatAnswerGenerationRequest):
    try:
        prompt = ChatPromptTemplate.from_messages([
            ("system", """당신은 지원자의 포트폴리오 데이터를 분석하여 채용 담당자의 예상 질문에 대한 핵심 답변 초안을 작성하는 전문가입니다.

[작성 지침]
1. 반드시 제공된 '포트폴리오 컨텍스트'에 실시간으로 존재하는 프로젝트와 정보만 사용하세요.
2. 과거에 있었으나 현재 컨텍스트에서 사라진 프로젝트에 대해서는 절대 언급하지 마세요. (매우 중요)
3. 지원자가 직접 말하는 것처럼 1인칭 시점('-했습니다', '-입니다')으로 작성하세요.
4. 각 답변은 3-4문장 이내로 명확하고 설득력 있게 작성하세요.
5. 마크다운 형식이나 이모지(Emoji)를 절대 사용하지 말고 순수 텍스트로만 작성하세요.
6. 반드시 아래 JSON 형식으로만 반환하세요.
{{
  "core_skills": "질문 1에 대한 답변",
  "main_stack": "질문 2에 대한 답변",
  "tech_depth": "질문 3에 대한 답변",
  "documentation": "질문 4에 대한 답변",
  "role_contribution": "질문 5에 대한 답변",
  "collaboration": "질문 6에 대한 답변",
  "cycle": "질문 7에 대한 답변",
  "artifacts": "질문 8에 대한 답변",
  "best_project": "질문 9에 대한 답변",
  "troubleshooting": "질문 10에 대한 답변",
  "decision_making": "질문 11에 대한 답변",
  "quantitative_performance": "질문 12에 대한 답변"
}}
"""),
            ("human", """다음 질문들에 대해 지원자의 입장에서 전문적인 답변 초안을 작성해주세요:
[1. 핵심 역량 및 기술 요약]
1-1. 지원자의 핵심 역량 3가지를 요약한다면?
1-2. 이 포트폴리오에서 가장 주력으로 사용한 '기술 스택(Main Skill)'은 무엇인가요?
1-3. 기술적으로 가장 깊이 있게 파고들거나 연구해 본 분야는 어디인가요?
1-4. 코드 작성 외에 설계 문서(API 명세, 기획서 등)도 작성할 줄 아나요?

[2. 역할 및 기여도 검증]
2-1. 각 프로젝트에서의 지원자의 구체적인 역할과 기여도는 어땠나요?
2-2. 팀 프로젝트에서 동료들과의 협업(코드 리뷰, 일정 관리)은 어떻게 진행했나요?
2-3. 기획부터 배포/운영까지 '전체 사이클'을 경험해 본 프로젝트가 있나요?
2-4. 실제 작성한 소스 코드나 디자인 원본 파일(Figma 등)을 볼 수 있나요?

[3. 문제 해결 및 성과]
3-1. 포트폴리오 중 가장 자신 있는 프로젝트 하나를 소개한다면?
3-2. 개발(또는 진행) 중 발생한 가장 치명적인 문제와 해결 과정은 무엇인가요?
3-3. 해당 기술(또는 디자인 컨셉)을 선정하게 된 특별한 이유나 논리가 있나요?
3-4. 프로젝트를 통해 얻은 구체적인 수치 성과(사용자 수, 성능 개선율 등)가 있나요?

포트폴리오 데이터:
{input}""")
        ])
        chain = prompt | llm
        response = chain.invoke({"input": request.portfolio_context})
        
        content = extract_text_from_response(response)
        print(f"DEBUG: Raw AI Response -> {content}") # 디버깅용 로그

        # JSON 추출 시도 (여러 패턴 고려)
        json_content = None
        
        # 1. ```json ... ``` 패턴
        json_match = re.search(r'```json\s*(\{.*?\})\s*```', content, re.DOTALL)
        if json_match:
            json_content = json_match.group(1)
        else:
            # 2. ``` ... ``` 패턴
            json_match = re.search(r'```\s*(\{.*?\})\s*```', content, re.DOTALL)
            if json_match:
                json_content = json_match.group(1)
            else:
                # 3. { ... } 중괄호 패턴
                json_match = re.search(r'(\{.*\})', content, re.DOTALL)
                if json_match:
                    json_content = json_match.group(1)

        if json_content:
            try:
                data = json.loads(json_content.strip())
                # 필수 키 검증
                required_keys = ["best_project", "role_contribution", "core_skills"]
                for key in required_keys:
                    if key not in data:
                        data[key] = "정보를 바탕으로 답변을 작성하지 못했습니다. 직접 입력해 주세요."
                return data
            except json.JSONDecodeError as je:
                print(f"❌ JSON 파싱 에러: {je}\nContent: {json_content}")
                return {
                    "error": f"JSON 형식이 올바르지 않습니다: {str(je)}",
                    "raw_content": content
                }
        else:
            print(f"❌ JSON 패턴을 찾을 수 없음: {content}")
            return {
                "error": "AI 응답에서 JSON 데이터를 찾을 수 없습니다.",
                "raw_content": content
            }
            
    except Exception as e:
        print(f"❌ 답변 생성 실패: {e}")
        return {"error": str(e)}

# DB 세션
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# --- [API] 포트폴리오 저장 ---
@app.post("/save-portfolio")
def save_portfolio(data: PortfolioUpdate, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.email == data.email).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    user.portfolio_data = json.dumps(data.portfolio_data)
    db.commit()
    return {"message": "Portfolio saved successfully"}

# --- [API] 포트폴리오 불러오기 ---
@app.get("/get-portfolio/{email}")
def get_portfolio(email: str, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.email == email).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    if not user.portfolio_data:
        raise HTTPException(status_code=404, detail="Portfolio data not found")

    return {"portfolio_data": json.loads(user.portfolio_data)}



# --- [API 1] 이메일 회원가입 ---
@app.post("/signup")
def signup(user: UserCreate, db: Session = Depends(get_db)):
    existing_user = db.query(User).filter(User.email == user.email).first()
    if existing_user:
        raise HTTPException(status_code=400, detail="이미 등록된 이메일입니다.")
    
    hashed_password = pwd_context.hash(user.password)
    new_user = User(email=user.email, password=hashed_password, name=user.name)
    db.add(new_user)
    db.commit()
    return {"message": "회원가입 성공"}

# --- [API 2] 이메일 로그인 ---
@app.post("/login")
def login(user: UserLogin, db: Session = Depends(get_db)):
    db_user = db.query(User).filter(User.email == user.email).first()
    if not db_user or not pwd_context.verify(user.password, db_user.password):
        raise HTTPException(status_code=400, detail="이메일 또는 비밀번호가 틀렸습니다.")
    
    portfolio_data = json.loads(db_user.portfolio_data) if db_user.portfolio_data else None
    return {"message": "로그인 성공", "user_name": db_user.name, "email": db_user.email, "portfolio_data": portfolio_data}

# --- [API 3] 구글 로그인 ---
@app.post("/google-login")
def google_login(data: GoogleToken, db: Session = Depends(get_db)):
    try:
        id_info = id_token.verify_oauth2_token(data.token, google_requests.Request())
        email = id_info['email']
        name = id_info.get('name', 'Google User')

        db_user = db.query(User).filter(User.email == email).first()
        if not db_user:
            new_user = User(email=email, password="SOCIAL_GOOGLE", name=name)
            db.add(new_user)
            db.commit()
            db_user = new_user
        
        portfolio_data = json.loads(db_user.portfolio_data) if db_user.portfolio_data else None
        return {"message": "구글 로그인 성공", "user_name": db_user.name, "email": db_user.email, "portfolio_data": portfolio_data}
    except ValueError:
        raise HTTPException(status_code=400, detail="유효하지 않은 구글 토큰입니다.")

# --- [API 4] 카카오 로그인 ---
@app.post("/kakao-login")
def kakao_login(data: KakaoToken, db: Session = Depends(get_db)):
    try:
        headers = {'Authorization': f'Bearer {data.token}'}
        me_res = requests.get("https://kapi.kakao.com/v2/user/me", headers=headers)
        me_data = me_res.json()
        
        kakao_account = me_data.get('kakao_account')
        if not kakao_account:
             raise HTTPException(status_code=400, detail="카카오 계정 정보를 불러올 수 없습니다.")

        email = kakao_account.get('email')
        profile = kakao_account.get('profile')
        nickname = profile.get('nickname') if profile else 'Kakao User'
        
        # 이메일 동의 안 했을 경우 임시 아이디 생성
        if not email:
             email = f"{me_data['id']}@kakao.temp" 

        db_user = db.query(User).filter(User.email == email).first()
        if not db_user:
            new_user = User(email=email, password="SOCIAL_KAKAO", name=nickname)
            db.add(new_user)
            db.commit()
            db_user = new_user
            
        portfolio_data = json.loads(db_user.portfolio_data) if db_user.portfolio_data else None
        return {"message": "카카오 로그인 성공", "user_name": db_user.name, "email": db_user.email, "portfolio_data": portfolio_data}
    except Exception as e:
        print("카카오 에러:", e)
        raise HTTPException(status_code=400, detail="카카오 로그인 실패")

# --- [API 5] 네이버 로그인 (추가됨) ---
@app.post("/naver-login")
def naver_login(data: NaverToken, db: Session = Depends(get_db)):
    try:
        # 네이버에 토큰 확인 요청
        headers = {'Authorization': f'Bearer {data.token}'}
        res = requests.get("https://openapi.naver.com/v1/nid/me", headers=headers)
        info = res.json()
        
        if info.get('resultcode') != '00':
            raise Exception("네이버 인증 실패")

        naver_account = info['response']
        email = naver_account.get('email')
        name = naver_account.get('name', 'Naver User')

        if not email:
             raise HTTPException(status_code=400, detail="이메일 정보가 없습니다.")

        # DB 확인 및 가입
        db_user = db.query(User).filter(User.email == email).first()
        if not db_user:
            new_user = User(email=email, password="SOCIAL_NAVER", name=name)
            db.add(new_user)
            db.commit()
            db_user = new_user
            
        portfolio_data = json.loads(db_user.portfolio_data) if db_user.portfolio_data else None
        return {"message": "네이버 로그인 성공", "user_name": db_user.name, "email": db_user.email, "portfolio_data": portfolio_data}
        
    except Exception as e:
        print("네이버 에러:", e)
        raise HTTPException(status_code=400, detail="네이버 로그인 실패")

# --- [API 6] AI 포트폴리오 생성 ---

portfolio_prompt = ChatPromptTemplate.from_messages([
    ("system", """
    당신은 전문 웹 디자이너입니다. 사용자 정보를 바탕으로 포트폴리오 웹사이트 JSON 데이터를 생성하세요.
    Markdown 코드블럭 없이 순수 JSON 문자열만 출력하세요.
    {{
        "theme": {{ "color": "#HEX", "font": "sans", "mood_emoji": "🚀", "layout": "gallery_grid" }},
        "hero": {{ "title": "제목", "subtitle": "부제", "tags": ["태그"] }},
        "about": {{ "intro": "소개", "description": "내용" }},
        "projects": [ {{ "title": "제목", "desc": "설명", "detail": "상세", "tags": ["기술"] }} ],
        "contact": {{ "email": "이메일", "github": "링크" }}
    }}
    """),
    ("human", "{input}")
])
portfolio_chain = portfolio_prompt | llm

@app.post("/submit")
def submit_data(data: UserAnswers):
    print("📢 [생성 요청] AI 작업 시작...")
    log_ai_usage(prompt_type="auto_generate")
    answers = data.answers
    projects_str = ""
    
    # 직무와 상관없이 최대 6개 프로젝트까지 포함
    for i in range(1, 7):
        # 디자이너용 필드와 일반용 필드 모두 확인
        title = answers.get(f"project{i}_title") or answers.get(f"design_project{i}_title")
        if title: projects_str += f"- 프로젝트 {i}: {title}\n"

    try:
        result = portfolio_chain.invoke({
            "input": f"이름:{answers.get('name')} 직무:{answers.get('job')} 강점:{answers.get('strength')} 분위기:{answers.get('moods')} 경력:{answers.get('career_summary')} 프로젝트:{projects_str}"
        })
        
        # JSON 정제
        content = result.content.replace("```json", "").replace("```", "").strip()
        match = re.search(r'\{.*\}', content, re.DOTALL)
        if match: content = match.group(0)
        
        return {"status": "success", "message": "완료!", "data": json.loads(content)}
    except Exception as e:
        print(f"❌ 생성 실패: {e}")
        return {"status": "error", "message": str(e)}

# --- [API 6.5] 이력서 파싱/분석 (새로 추가됨) ---
class ResumeAnalyzeRequest(BaseModel):
    resumeText: str
    images: list[str] = []

def extract_text_from_pdf(file_bytes):
    import pypdf
    import io
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = pypdf.PdfReader(pdf_file)
        text = ""
        total_pages = len(reader.pages)
        print(f"📖 PDF 총 페이지 수: {total_pages}")
        
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text += page_text + "\n"
                    print(f"  ✅ 페이지 {page_num + 1}: {len(page_text)} 글자 추출")
                else:
                    print(f"  ⚠️ 페이지 {page_num + 1}: 텍스트 없음 (이미지 전용 페이지일 수 있음)")
            except Exception as e:
                print(f"  ❌ 페이지 {page_num + 1} 추출 실패: {e}")
                continue
        
        if not text.strip():
            print("⚠️ PDF에서 텍스트를 추출할 수 없습니다. 스캔된 이미지 PDF이거나 보호된 문서일 수 있습니다.")
            return ""
        
        print(f"✅ PDF 파싱 성공: 총 {len(text)} 글자 추출")
        return text.strip()
    except Exception as e:
        print(f"❌ PDF 파싱 오류: {e}")
        import traceback
        traceback.print_exc()
        raise


def extract_images_from_pdf(file_bytes):
    import fitz  # PyMuPDF
    import base64
    
    images = []
    try:
        # PDF 문서 열기
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        
        # 최대 5페이지만 처리 (토큰 및 시간 절약)
        max_pages = min(len(doc), 5)
        print(f"🖼️ PDF 렌더링 시작 (총 {len(doc)}페이지 중 {max_pages}페이지만 처리)")
        
        for i in range(max_pages):
            page = doc.load_page(i)
            # 페이지를 이미지(Pixmap)로 렌더링 (해상도 조절 가능, 기본 72dpi -> matrix로 확대 가능)
            # matrix=fitz.Matrix(2, 2) -> 2배 확대 (약 144dpi) - 글자 가독성 위해 권장
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            
            # 이미지 바이트 추출 (PNG)
            img_bytes = pix.tobytes("png")
            encoded = base64.b64encode(img_bytes).decode('utf-8')
            
            # data URL 형식으로 변환
            data_url = f"data:image/png;base64,{encoded}"
            images.append(data_url)
            print(f"  ✅ P{i+1} 렌더링 완료 ({len(img_bytes)} bytes)")
            
        return images
    except Exception as e:
        print(f"❌ PDF 렌더링/이미지 추출 오류: {e}")
        import traceback
        traceback.print_exc()
        return []


def extract_images_from_docx(file_bytes):
    """DOCX 파일에서 이미지를 추출하여 base64 인코딩된 데이터 URL 리스트로 반환"""
    import docx
    import io
    import base64
    
    images = []
    try:
        doc_file = io.BytesIO(file_bytes)
        doc = docx.Document(doc_file)
        
        # 문서 내 모든 관계(relationships)에서 이미지 찾기
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    # 이미지 타입 감지
                    content_type = rel.target_part.content_type
                    # base64 인코딩
                    encoded = base64.b64encode(image_data).decode('utf-8')
                    data_url = f"data:{content_type};base64,{encoded}"
                    images.append(data_url)
                    print(f"  📷 이미지 추출: {content_type}, {len(image_data)} bytes")
                except Exception as e:
                    print(f"  ⚠️ 이미지 추출 실패: {e}")
                    continue
        
        print(f"✅ DOCX 이미지 추출 완료: {len(images)}개")
        return images
    except Exception as e:
        print(f"❌ DOCX 이미지 추출 오류: {e}")
        return []

def extract_text_from_docx(file_bytes):
    import docx
    import io
    try:
        doc_file = io.BytesIO(file_bytes)
        doc = docx.Document(doc_file)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        text = "\n".join(paragraphs)
        print(f"✅ DOCX 파싱 성공: {len(paragraphs)} 문단, {len(text)} 글자 추출")
        return text.strip()
    except Exception as e:
        print(f"❌ DOCX 파싱 오류: {e}")
        raise

@app.post("/api/parse-resume")
async def parse_resume(file: UploadFile = File(...)):
    try:
        print(f"📄 파일 업로드 시작: {file.filename} ({file.content_type})")
        contents = await file.read()
        print(f"📦 파일 크기: {len(contents)} bytes")
        
        filename = file.filename.lower()
        extracted_text = ""
        extracted_images = []

        if filename.endswith(".pdf"):
            print("🔍 PDF 파싱 시작...")
            extracted_text = extract_text_from_pdf(contents)
            print("🔍 PDF 이미지 추출 시작...")
            extracted_images = extract_images_from_pdf(contents)
        elif filename.endswith(".docx"):
            print("🔍 DOCX 파싱 시작...")
            extracted_text = extract_text_from_docx(contents)
            print("🔍 DOCX 이미지 추출 시작...")
            extracted_images = extract_images_from_docx(contents)
        elif filename.endswith(".txt"):
            print("🔍 TXT 파싱 시작...")
            extracted_text = contents.decode("utf-8")
            extracted_images = []
        else:
            print(f"❌ 지원하지 않는 파일 형식: {filename}")
            return {"error": "지원하지 않는 파일 형식입니다. (PDF, DOCX, TXT 지원)"}

        if not extracted_text or len(extracted_text.strip()) == 0:
            print("⚠️ 경고: 추출된 텍스트가 비어있습니다!")
            return {"error": "파일에서 텍스트를 추출할 수 없습니다. 파일이 비어있거나 이미지만 포함되어 있을 수 있습니다."}
        
        print(f"✅ 파싱 완료: {len(extracted_text)} 글자, {len(extracted_images)} 이미지")
        return {"text": extracted_text, "filename": file.filename, "images": extracted_images}

    except Exception as e:
        print(f"❌ 파일 파싱 실패: {e}")
        import traceback
        traceback.print_exc()
        return {"error": f"파일 파싱 중 오류가 발생했습니다: {str(e)}"}

@app.post("/api/analyze-resume")
def analyze_resume(request: ResumeAnalyzeRequest):
    try:
        log_ai_usage(prompt_type="resume_analysis")
        
        if request.images:
            print(f"🖼️ 이미지 분석 모드: {len(request.images)}개의 이미지 포함")
            
            message_content = []
            
            # 시스템 프롬프트 내용을 텍스트로 추가
            system_prompt = """당신은 채용 전문가 AI입니다. 제공된 이력서 이미지와 텍스트를 종합적으로 분석하여 구조화된 JSON 데이터로 변환해주세요.
            
            [분석 요구사항]
            1. 이름, 연락처, 이메일 등 기본 정보를 추출하세요.
            2. 핵심 기술(Skills)을 리스트로 추출하세요.
            3. 경력 사항을 요약하여 'career_summary'에 작성하세요 (예: "총 5년차, 주요 경력: ABC사, XYZ사").
            4. 이력서에 명시된 '모든' 주요 프로젝트 경험을 요약하여 'projects' 배열에 담으세요 (개수 제한 없음).
            5. 자기소개나 포트폴리오에 쓸만한 문구를 'intro'에 작성하세요.
            
            [출력 포맷 (JSON Only)]
            {{
                "name": "지원자 이름",
                "phone": "010-XXXX-XXXX",
                "email": "email@example.com",
                "link": "github/blog url",
                "intro": "한줄 소개",
                "career_summary": "경력 요약 텍스트",
                "skills": ["Skill1", "Skill2", "Skill3"],
                "projects": [
                    {{ "title": "프로젝트명", "desc": "프로젝트 설명 및 역할", "duration": "기간" }}
                ]
            }}
            """
            
            # 텍스트가 있으면 추가
            user_input = "다음 이력서(이미지 포함)를 분석해주세요."
            if request.resumeText:
                user_input += f"\n\n[추출된 텍스트]\n{request.resumeText}"
                
            message_content.append({"type": "text", "text": system_prompt + "\n\n" + user_input})
            
            # 이미지들 추가
            for img_data in request.images:
                # data:image/jpeg;base64,... 형식 파싱
                if "," in img_data:
                    header, base64_str = img_data.split(",", 1)
                    # image_url 방식을 사용 (langchain_google_genai 지원 방식)
                    message_content.append({
                        "type": "image_url", 
                        "image_url": {"url": img_data}
                    })
                else:
                    # 헤더가 없는 경우 처리하지 않거나 기본값 가정
                    pass
            
            msg = HumanMessage(content=message_content)
            response = llm.invoke([msg])
            
        else:
            # 텍스트 전용 모드 (기존 로직)
            prompt = ChatPromptTemplate.from_messages([
                ("system", """당신은 채용 전문가 AI입니다. 이력서 텍스트를 분석하여 구조화된 JSON 데이터로 변환해주세요.
                
                [분석 요구사항]
                1. 이름, 연락처, 이메일 등 기본 정보를 추출하세요.
                2. 핵심 기술(Skills)을 리스트로 추출하세요.
                3. 경력 사항을 요약하여 'career_summary'에 작성하세요 (예: "총 5년차, 주요 경력: ABC사, XYZ사").
                4. 이력서에 명시된 '모든' 주요 프로젝트 경험을 요약하여 'projects' 배열에 담으세요 (개수 제한 없음).
                5. 자기소개나 포트폴리오에 쓸만한 문구를 'intro'에 작성하세요.
                
                [출력 포맷 (JSON Only)]
                {{
                    "name": "지원자 이름",
                    "phone": "010-XXXX-XXXX",
                    "email": "email@example.com",
                    "link": "github/blog url",
                    "intro": "한줄 소개",
                    "career_summary": "경력 요약 텍스트",
                    "skills": ["Skill1", "Skill2", "Skill3"],
                    "projects": [
                        {{ "title": "프로젝트명", "desc": "프로젝트 설명 및 역할", "duration": "기간" }}
                    ]
                }}
                """),
                ("human", "다음 이력서 내용을 분석해주세요:\n\n{input}")
            ])
            
            chain = prompt | llm
            response = chain.invoke({"input": request.resumeText})
        
        # JSON 추출 - response.content가 리스트일 수 있으므로 먼저 텍스트로 변환
        content = extract_text_from_response(response)
        print(f"🤖 AI 응답 길이: {len(content)} 글자")
        
        json_match = re.search(r'```json\s*(\{.*?\})\s*```', content, re.DOTALL)
        if json_match:
            json_content = json_match.group(1)
        else:
            json_match = re.search(r'(\{.*\})', content, re.DOTALL)
            json_content = json_match.group(0) if json_match else "{}"

        parsed_data = json.loads(json_content)
        print(f"✅ 이력서 분석 완료: {parsed_data.get('name', 'Unknown')}")
        return parsed_data

    except Exception as e:
        print(f"❌ 이력서 분석 실패: {e}")
        import traceback
        traceback.print_exc()
        return {"error": str(e)}

# --- [API 7] 챗봇 ---
def extract_text_from_response(response):
    """
    Gemini API 응답에서 실제 텍스트만 추출하는 헬퍼 함수
    응답이 문자열이면 그대로 반환, 객체/리스트면 텍스트 추출
    """
    content = response.content
    
    # 이미 문자열이면 그대로 반환
    if isinstance(content, str):
        return content
    
    # 리스트인 경우 (예: [{'type': 'text', 'text': '...', 'extras': {...}}])
    if isinstance(content, list):
        text_parts = []
        for item in content:
            if isinstance(item, dict) and 'text' in item:
                text_parts.append(item['text'])
            elif isinstance(item, str):
                text_parts.append(item)
        return ''.join(text_parts)
    
    # 딕셔너리인 경우
    if isinstance(content, dict):
        if 'text' in content:
            return content['text']
    
    # 그 외의 경우 문자열로 변환
    return str(content)

@app.post("/chat")
def chat_bot(request: ChatRequest):
    try:
        # 1. 포포(Popo) 모드: 포트폴리오 제작 도우미
        if not request.is_shared:
            popo_prompt = ChatPromptTemplate.from_messages([
                ("system", """당신은 친절하고 전문적인 포트폴리오 코치 '포포(Popo)'입니다.
사용자가 자신의 강점을 잘 드러내는 포트폴리오를 완성할 수 있도록 돕는 것이 당신의 역할입니다.

[상담 지침]
1. 사용자가 입력한 현재 포트폴리오 정보(context)가 있다면 이를 분석하여 개선점을 제안하세요.
2. 구체적인 피드백을 제공하되, 격려와 응원을 아끼지 마세요.
3. 포트폴리오 구성, 직무별 핵심 역량 강조 방법, 프로젝트 요약 기술 등에 대해 조언하세요.
4. 사용자 정보에 기반하여 답변하되, 부족한 부분은 질문을 통해 보완할 수 있게 유도하세요.

{context}
"""),
                ("human", "{input}")
            ])
            
            context_str = request.portfolio_context if request.portfolio_context else "아직 입력된 포트폴리오 정보가 없습니다."
            chat_chain = popo_prompt | llm
            
            # Log usage
            log_ai_usage(prompt_type="popo")
            
            response = chat_chain.invoke({
                "input": request.message,
                "context": f"현재 포트폴리오 정보: {context_str}"
            })

        # 2. 무무(Mumu) 모드: 포트폴리오 도슨트 (인사담당자 대응)
        else:
            mumu_prompt = ChatPromptTemplate.from_messages([
                ("system", """당신은 지원자의 포트폴리오를 전문적으로 설명하고 안내하는 '도슨트 무무'입니다.
인사담당자(채용 담당자)에게 지원자의 역량을 신뢰감 있게 전달하는 것이 당신의 목표입니다.

당신은 지원자의 포트폴리오를 전문적으로 설명하는 '도슨트 무무'입니다.
지원자를 대신하여 채용 담당자에게 신뢰감 있는 정보를 전달하는 역할을 수행합니다.

[핵심 원칙]
1. '지원자가 직접 검수한 정보(Verified)'가 있다면 이를 최우선으로 활용하여 답변하세요. 이 경우 "지원자가 직접 확인한 정보에 따르면"과 같은 문구를 포함하세요.
2. 직접 입력된 답변이 없는 질문의 경우, '포트폴리오 데이터'에 기반하여 객관적인 사실만 요약해서 전달하세요.
3. **절대 '추측'하거나 '생각됩니다'와 같은 불확실한 표현을 사용하지 마세요.** (매우 중요)
4. 대신 "기재된 프로젝트 기록을 분석한 바로는...", "등록된 기술 스택에 따르면..."과 같이 데이터에 근거한 확신 있는 말투를 사용하세요.
5. 만약 데이터 자체가 아예 없는 내용이라면 지어내지 말고, "해당 상세 내용은 현재 자료에서 확인되지 않습니다. 지원자분께 직접 문의하여 더 자세한 이야기를 들어보시는 것을 추천드립니다."라고 정중히 안내하세요.
6. 전문적이고 정중하며, 지원자를 높여주는 대리인으로서의 톤을 유지하세요."""),
                ("human", "{input}")
            ])
            
            context_str = request.portfolio_context if request.portfolio_context else "포트폴리오 정보가 제공되지 않았습니다."
            chat_chain = mumu_prompt | llm
            
            # Log usage
            log_ai_usage(prompt_type="mumu")
            
            response = chat_chain.invoke({
                "input": request.message,
                "context": f"사용자 상세 데이터: {context_str}"
            })
        
        # 응답에서 실제 텍스트만 추출
        reply_text = extract_text_from_response(response)
        return {"reply": reply_text}
    except Exception as e:
        print(f"❌ 챗봇 오류: {e}")
        import traceback
        traceback.print_exc()
        return {"reply": "죄송합니다. 응답 생성 중 오류가 발생했습니다."}
# ==================== ADMIN API (SUPABASE) ====================
from admin_apis import (
    get_admin_stats as admin_stats_handler,
    get_all_users as admin_users_handler,
    delete_user as admin_delete_user_handler,
    get_all_portfolios as admin_portfolios_handler
)
from admin_auth import verify_admin

@app.get('/api/admin/stats')
def admin_stats_route(admin_email: str = Depends(verify_admin)):
    return admin_stats_handler(admin_email)

@app.get('/api/admin/users')
def admin_users_route(skip: int = 0, limit: int = 50, search: str = None, admin_email: str = Depends(verify_admin)):
    return admin_users_handler(skip, limit, search, admin_email)

@app.delete('/api/admin/users/{user_id}')
def admin_delete_user_route(user_id: str, admin_email: str = Depends(verify_admin)):
    return admin_delete_user_handler(user_id, admin_email)

@app.get('/api/admin/portfolios')
def admin_portfolios_route(skip: int = 0, limit: int = 50, search: str = None, admin_email: str = Depends(verify_admin)):
    return admin_portfolios_handler(skip, limit, search, admin_email)

from admin_apis import batch_delete_users as admin_batch_delete_users_handler
from pydantic import BaseModel

class BatchDeleteRequest(BaseModel):
    user_ids: list[str]

@app.post('/api/admin/users/batch-delete')
def admin_batch_delete_users_route(request: BatchDeleteRequest, admin_email: str = Depends(verify_admin)):
    return admin_batch_delete_users_handler(request.user_ids, admin_email)


# --- 새로운 관리 기능 (Notices, AI Stats, Template Config) ---
from admin_apis import (
    get_notices, get_active_notices, create_notice, update_notice, delete_notice,
    NoticeCreate, NoticeUpdate,
    get_ai_stats,
    get_template_configs, update_template_config, TemplateConfigUpdate,
    log_ai_usage
)

# 1. 공지사항 라우트
@app.get('/api/notices/active')
def get_active_notices_route():
    return get_active_notices()

@app.get('/api/admin/notices')
def admin_get_notices(skip: int = 0, limit: int = 20, admin_email: str = Depends(verify_admin)):
    return get_notices(skip, limit, admin_email)

@app.post('/api/admin/notices')
def admin_create_notice(notice: NoticeCreate, admin_email: str = Depends(verify_admin)):
    return create_notice(notice, admin_email)

@app.put('/api/admin/notices/{notice_id}')
def admin_update_notice(notice_id: str, notice: NoticeUpdate, admin_email: str = Depends(verify_admin)):
    return update_notice(notice_id, notice, admin_email)

@app.delete('/api/admin/notices/{notice_id}')
def admin_delete_notice(notice_id: str, admin_email: str = Depends(verify_admin)):
    return delete_notice(notice_id, admin_email)

# 2. AI 통계 라우트
@app.get('/api/admin/stats/ai')
def admin_get_ai_stats(period: str = 'daily', admin_email: str = Depends(verify_admin)):
    return get_ai_stats(period, admin_email)


# 3. 템플릿 설정 라우트
# Public endpoint for reading template config (no auth required)
@app.get('/api/templates/config')
def public_get_template_configs():
    return get_template_configs(admin_email=None)

# Admin endpoint for reading template config (auth required)
@app.get('/api/admin/templates/config')
def admin_get_template_configs(admin_email: str = Depends(verify_admin)):
    return get_template_configs(admin_email=admin_email)

@app.put('/api/admin/templates/config/{key}')
def admin_update_template_config(key: str, config: TemplateConfigUpdate, admin_email: str = Depends(verify_admin)):
    return update_template_config(key, config, admin_email)

